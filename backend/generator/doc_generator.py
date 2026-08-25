from docx import Document
import os

def generate_compliance_sheet(bid_id: str, company_name: str, output_dir: str = "./outputs") -> str:
    """
    Generates customized technical compliance sheets and affidavits using python-docx.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    
    # Document Header
    doc.add_heading(f"Technical Compliance Matrix - Bid #{bid_id}", level=1)
    doc.add_paragraph(f"Prepared for: {company_name}")
    doc.add_paragraph("This document certifies that all parameters meet the NIT specifications and complies with all mandatory portal criteria.")
    
    # Specification Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Requirement Clause"
    hdr_cells[1].text = "Specified Standard"
    hdr_cells[2].text = "Bidder Compliance Status"
    
    # Sample Row
    row_cells = table.add_row().cells
    row_cells[0].text = "Annual Turnover"
    row_cells[1].text = "Min ₹25,00,000"
    row_cells[2].text = "Complied (Certified by CA with UDIN)"
    
    # Affidavits Section
    doc.add_heading("Non-Blacklisting Declaration", level=2)
    doc.add_paragraph("I/We hereby declare that our firm has not been blacklisted or debarred by any Central/State Government department, PSU, or autonomous body.")
    
    # Sanitize bid_id for safe filenames
    safe_bid_id = bid_id.replace("/", "_")
    file_path = os.path.join(output_dir, f"Compliance_{safe_bid_id}.docx")
    doc.save(file_path)
    return file_path
